from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import sys


def ensure_dep(name: str):
    try:
        __import__(name)
    except Exception as e:
        raise SystemExit(
            f"Missing dependency '{name}'. Install it with:\n\n"
            f"  python -m pip install {name}\n\n"
            f"Original error: {e}"
        )


ensure_dep("docx")
from docx import Document  # type: ignore


@dataclass
class Line:
    kind: str  # heading1|heading2|heading3|bullet|text|hr|blank
    text: str


def parse_md(md: str) -> list[Line]:
    out: list[Line] = []
    for raw in md.splitlines():
        line = raw.rstrip("\n")

        if not line.strip():
            out.append(Line("blank", ""))
            continue

        if line.strip() == "---":
            out.append(Line("hr", ""))
            continue

        if line.startswith("### "):
            out.append(Line("heading3", line[4:].strip()))
            continue

        if line.startswith("## "):
            out.append(Line("heading2", line[3:].strip()))
            continue

        if line.startswith("# "):
            out.append(Line("heading1", line[2:].strip()))
            continue

        m = re.match(r"^\s*-\s+(.*)$", line)
        if m:
            out.append(Line("bullet", m.group(1).strip()))
            continue

        out.append(Line("text", line.strip()))

    return out


def add_paragraph_with_bold_markers(doc: Document, text: str, style: str | None = None):
    """
    Supports **bold** markers inside a line.
    """
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def render_docx(lines: list[Line]) -> Document:
    doc = Document()

    for ln in lines:
        if ln.kind == "blank":
            doc.add_paragraph("")
            continue

        if ln.kind == "hr":
            doc.add_paragraph("—" * 30)
            continue

        if ln.kind == "heading1":
            doc.add_heading(ln.text, level=1)
            continue

        if ln.kind == "heading2":
            doc.add_heading(ln.text, level=2)
            continue

        if ln.kind == "heading3":
            doc.add_heading(ln.text, level=3)
            continue

        if ln.kind == "bullet":
            add_paragraph_with_bold_markers(doc, ln.text, style="List Bullet")
            continue

        add_paragraph_with_bold_markers(doc, ln.text)

    return doc


def main() -> int:
    repo = Path(__file__).resolve().parents[1]
    md_path = repo / "docs" / "manuel-contentful-ecd.md"
    out_path = repo / "docs" / "Manuel-Contentful-ECD.docx"

    if not md_path.exists():
        print(f"Markdown source not found: {md_path}", file=sys.stderr)
        return 1

    md = md_path.read_text(encoding="utf-8")
    lines = parse_md(md)
    doc = render_docx(lines)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

